import hashlib
import importlib.util
from pathlib import Path

import pytest
from docx import Document

from omni_healthcheck.reporting import ReportModel
from omni_healthcheck.v4_adapter import _prepare_unit, build_v4_report
from omni_healthcheck.v4_quality import V4QualityError, validate_v4_report
from omni_healthcheck.v4_renderer import VENDOR_RENDERER, _vendor_renderer
from omni_healthcheck.docx_renderer import _font_config_for_platform


ROOT = Path(__file__).parents[1]


def test_linux_pdf_conversion_uses_system_fontconfig() -> None:
    assert _font_config_for_platform("linux") is None


def test_macos_pdf_conversion_uses_project_fallback() -> None:
    path = _font_config_for_platform("darwin")
    assert path is not None
    assert path.name == "fonts.macos.conf"


def _model() -> ReportModel:
    return ReportModel(
        customer="測試客戶",
        system_name="db-system",
        period="2026-H1",
        engineer="XXX",
        product="EPAS",
        nodes=[
            {"hostname": "db-primary", "role": "Primary", "services": ["EFM"]},
            {"hostname": "db-standby", "role": "Standby", "services": ["EFM"]},
            {"hostname": "db-dr", "role": "DR", "services": []},
            {"hostname": "pem-server", "role": "Witness", "services": ["PEM"]},
            {"hostname": "efm-witness", "role": "Witness", "services": ["EFM"]},
        ],
        summary={"normal": 1, "attention": 1, "critical": 0, "pending": 0},
        sections=[
            {
                "section_id": "3",
                "title": "作業系統健檢",
                "groups": [
                    {
                        "title": "3.1 主機與組態設定",
                        "units": [
                            {
                                "title": "主機彙整",
                                "headers": ["項目", "db-primary\nPrimary", "db-standby\nStandby"],
                                "rows": [
                                    ["作業系統", "RHEL 8", "RHEL 8"],
                                    ["CPU Core 數", "8", "8"],
                                    ["記憶體", "64 GB", "64 GB"],
                                ],
                                "omitted_rows": 0,
                                "assessment": None,
                            }
                        ],
                    }
                ],
            },
            {
                "section_id": "4",
                "title": "PostgreSQL 資料庫健檢",
                "groups": [
                    {
                        "title": "4.1 組態與安全設定",
                        "units": [
                            {
                                "title": "版本資訊",
                                "headers": ["Metric", "Value"],
                                "rows": [["Database Version", "PostgreSQL 16.14"]],
                                "omitted_rows": 0,
                                "assessment": None,
                            },
                            {
                                "title": "Primary postgresql.auto.conf",
                                "headers": ["Output"],
                                "rows": [["shared_buffers = '8GB'"]],
                                "omitted_rows": 0,
                                "assessment": {
                                    "status": "attention",
                                    "status_label": "注意",
                                    "observation": "跨節點比較有 1 項差異\n結論：本項需注意",
                                    "recommendation": "確認角色必要差異",
                                },
                            }
                        ],
                    }
                ],
            },
        ],
        findings=[],
        coverage={"coverage_percent": 100},
        cve={"status": "pending"},
    )


def test_adapter_preserves_primary_database_and_configuration_observation(
    tmp_path: Path,
) -> None:
    report = build_v4_report(
        _model(),
        {"evidence": []},
        tmp_path,
    )
    assert report["database_source_hostname"] == "db-primary"
    section_items = report["chapters"][1]["sections"][0]["items"]
    item = next(item for item in section_items if "postgresql.auto.conf" in item["title"])
    assert item["node"] == "db-primary"
    assert item["status"] == "注意"
    assert "跨節點比較" in item["observation"]
    assert item["controlled_continuation"] is True
    assert report["nodes"][0]["cpu"] == "8 cores"
    database_by_node = {
        node["hostname"]: node["database"] for node in report["nodes"]
    }
    assert database_by_node == {
        "db-primary": "EDB Postgres Advanced Server",
        "db-standby": "EDB Postgres Advanced Server",
        "db-dr": "EDB Postgres Advanced Server",
        "pem-server": "PostgreSQL",
        "efm-witness": "",
    }
    assert report["product"]["name"] == "EDB Postgres Advanced Server"
    assert report["cover_company_name"] == "Omniwaresoft Tech"
    assert report["show_components"] is False
    version = next(item for item in section_items if item["title"] == "版本資訊")
    assert version["evidence"]["rows"] == [
        ["Database Version", "EDB Postgres Advanced Server 16.14"]
    ]
    assert version["assessment_display"] is False
    assert "status" not in version
    assert "observation" not in version
    assert "recommendation" not in version
    assert validate_v4_report(report, {"evidence": []})["delivery_allowed"] is True


def test_adapter_preserves_ready_history_comparison(tmp_path: Path) -> None:
    model = _model().model_copy(update={"history": {
        "status": "ready",
        "prior_period": "2025-H2",
        "summary": {"improved": 1, "worsened": 2, "evidence_changed": 0, "added": 3, "removed": 1},
        "changes": [{
            "node": "db-primary", "check_id": "filesystem_usage",
            "change": "worsened", "prior_status": "normal",
            "current_status": "attention",
        }],
    }})
    report = build_v4_report(model, {"evidence": []}, tmp_path)
    assert report["history_comparison"]["status"] == "ready"
    assert report["history_comparison"]["summary"]["worsened"] == 2


def test_adapter_applies_requested_report_limits_and_database_columns() -> None:
    inventory = _prepare_unit(
        {
            "title": "資料庫清單",
            "headers": [
                "Name", "Owner", "Encoding", "Access privileges", "Size", "Description"
            ],
            "rows": [["app", "owner", "UTF8", "=Tc/owner", "10 GB", "omit me"]],
        }
    )
    assert inventory["headers"] == ["資料庫名稱", "擁有者", "權限", "大小"]
    assert inventory["rows"] == [["app", "owner", "=Tc/owner", "10 GB"]]

    txid = _prepare_unit(
        {
            "title": "Transaction ID 年齡",
            "headers": ["table", "txid_age"],
            "rows": [[f"t{value}", str(value)] for value in range(12)],
        }
    )
    assert len(txid["rows"]) == 10
    assert txid["rows"][0] == ["t11", "11"]

    indexes = _prepare_unit(
        {
            "title": "罕用索引",
            "headers": ["index", "idx_scan"],
            "rows": [["used", "5"], ["zero", "0"]] + [
                [f"i{value}", str(value)] for value in range(1, 12)
            ],
        }
    )
    assert len(indexes["rows"]) == 10
    assert indexes["rows"][0] == ["zero", "0"]

    replication = _prepare_unit(
        {
            "title": "同步狀態",
            "headers": [
                "pid", "usesysid", "usename", "client_hostname", "state", "sync_state"
            ],
            "rows": [["1", "2", "repuser", "standby", "streaming", "async"]],
        }
    )
    assert replication["headers"] == [
        "pid", "usename", "client_hostname", "state", "sync_state"
    ]
    assert replication["rows"] == [
        ["1", "repuser", "standby", "streaming", "async"]
    ]


def test_v4_quality_blocks_pending_scope() -> None:
    report = build_v4_report(_model(), {"evidence": []}, ROOT)
    ledger = {
        "evidence": [
            {"path": "unknown.png", "decision": "pending"}
        ]
    }
    with pytest.raises(V4QualityError, match="pending_evidence"):
        validate_v4_report(report, ledger)


def test_v4_quality_requires_authoritative_cve_fields() -> None:
    report = build_v4_report(_model(), {"evidence": []}, ROOT)
    report["version_updates"] = [{"current": "PostgreSQL 16.4", "recommended": "16.5", "summary": "測試", "cves": [{"id": "CVE-2026-1"}]}]
    result = validate_v4_report(report, {"evidence": []}, raise_on_failure=False)
    assert result["delivery_allowed"] is False
    assert any(item.startswith("cve.missing_authoritative_metadata") for item in result["failed_gates"])


def test_adapter_passes_deterministic_cve_cache_to_v4_contract(tmp_path: Path) -> None:
    model = _model().model_copy(update={"cve": {
        "status": "ready", "version_updates": [{
            "current": "EDB Postgres Advanced Server 15.7", "recommended": "15.8",
            "summary": "權威快取比對完成。", "cves": [{
                "id": "CVE-2026-0001", "summary": "測試", "cvss_score": "8.1",
                "severity": "HIGH", "cvss_version": "3.1", "vector": "CVSS:3.1/AV:N",
                "score_source": "https://vendor.example/advisory", "match_status": "applicable",
                "fixed_version": "15.8", "source": "https://vendor.example/advisory",
            }],
        }],
    }})
    report = build_v4_report(model, {"evidence": []}, tmp_path)
    assert report["version_updates"][0]["cves"][0]["match_status"] == "applicable"
    assert validate_v4_report(report, {"evidence": []})["delivery_allowed"] is True


def test_approved_v4_renderer_hash_is_pinned() -> None:
    digest = hashlib.sha256(VENDOR_RENDERER.read_bytes()).hexdigest()
    assert digest == "cff4dede7b4199d82659335dab453772f65d6255ea924ceef302ea669bb29c0d"


def test_renderer_finds_vendor_from_release_working_directory(tmp_path: Path, monkeypatch) -> None:
    renderer = tmp_path / "vendor/omni-v4-renderer/scripts/build_report.py"
    renderer.parent.mkdir(parents=True)
    renderer.write_text("# approved renderer fixture\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    monkeypatch.delenv("OMNICHECK_VENDOR_RENDERER", raising=False)
    assert _vendor_renderer() == renderer


def test_renderer_omits_assessment_block_for_information_only_item() -> None:
    spec = importlib.util.spec_from_file_location("omni_v4_info_only_test", VENDOR_RENDERER)
    assert spec is not None and spec.loader is not None
    renderer = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(renderer)
    document = Document()
    renderer.configure_document(document)
    renderer.add_item(document, {
        "title": "Extension 清單",
        "assessment_display": False,
        "evidence": {"type": "table", "headers": ["Extension"], "rows": [["pg_stat_statements"]]},
    })
    assert len(document.tables) == 1
    assert all("狀態" not in cell.text for table in document.tables for row in table.rows for cell in row.cells)


def test_v4_summary_rows_do_not_chain_the_whole_table_to_the_next_page() -> None:
    spec = importlib.util.spec_from_file_location("omni_v4_pagination_test", VENDOR_RENDERER)
    assert spec is not None and spec.loader is not None
    renderer = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(renderer)
    document = Document()
    renderer.configure_document(document)
    summary = [
        {
            "status": "注意",
            "item": f"檢查項目 {index}",
            "finding": "需要後續處理",
            "recommendation": "請安排確認",
            "reference": "4.2",
        }
        for index in range(20)
    ]

    renderer.add_updates_and_summary(document, {"summary": summary}, 6)

    table = document.tables[-1]
    assert len(table.rows) == 21
    assert all(
        paragraph.paragraph_format.keep_with_next
        for cell in table.rows[0].cells
        for paragraph in cell.paragraphs
    )
    assert all(
        not paragraph.paragraph_format.keep_with_next
        for row in table.rows[1:]
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
