"""Jiuxing V4-oriented deterministic DOCX renderer."""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from omni_healthcheck.reporting import ReportModel


FONT = "Arial Unicode MS"
ACCENT = "4F81BD"
TEAL = "00758A"
LABEL_BLUE = "C6D9F1"
OUTPUT_GRAY = "F2F2F2"
TEXT_GRAY = "808080"
STATUS_COLORS = {
    "normal": ("70AD47", "70AD47"),
    "attention": ("FF9933", "F28C28"),
    "critical": ("C00000", "C00000"),
    "pending": ("A6A6A6", "666666"),
}


def _font(run, size=10, bold=False, color="111111", name=FONT):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _remove_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def _repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def _cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _cell_text(cell, text, *, size=9.5, bold=False, color="111111", mono=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    _font(run, size=size, bold=bold, color=color, name="Menlo" if mono else FONT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_margins(cell)


def _set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Cm(width)


def _output_block(doc, headers, rows):
    if headers == ["Output"]:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_borders(table)
        text = "\n".join(row[0] for row in rows)
        _cell_text(table.cell(0, 0), text, size=8.2, mono=True)
        _shade(table.cell(0, 0), OUTPUT_GRAY)
        _set_widths(table, [17.2])
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_borders(table)
    for index, header in enumerate(headers):
        _cell_text(table.rows[0].cells[index], header, bold=True, size=8.5)
        _shade(table.rows[0].cells[index], LABEL_BLUE)
    _repeat_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        _cant_split(row)
        for index, value in enumerate(values):
            _cell_text(row.cells[index], value, size=8.3)
    _set_widths(table, [17.2 / len(headers)] * len(headers))


def _assessment_block(doc, assessment):
    status = assessment["status"]
    fill, color = STATUS_COLORS[status]
    container = doc.add_table(rows=1, cols=1)
    container.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_borders(container)
    _cant_split(container.rows[0])
    host = container.cell(0, 0)
    _cell_margins(host, top=0, start=0, bottom=0, end=0)
    table = host.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_borders(table)
    labels = ("狀態", "觀察", "建議")
    values = (
        assessment["status_label"],
        assessment["observation"],
        assessment["recommendation"],
    )
    for index, (label, value) in enumerate(zip(labels, values)):
        _cell_text(
            table.cell(index, 0), label,
            size=9.5, color="FFFFFF" if index == 0 else "111111",
        )
        _shade(table.cell(index, 0), fill if index == 0 else LABEL_BLUE)
        _cell_text(
            table.cell(index, 1), value,
            size=9.5, color=color if index == 0 else "111111",
        )
        if index == 0:
            _shade(table.cell(index, 1), OUTPUT_GRAY)
        _cant_split(table.rows[index])
    _set_widths(table, [4.0, 13.2])
    host.paragraphs[0]._element.getparent().remove(host.paragraphs[0]._element)


def _heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _font(run, size=16 if level == 1 else 12, bold=True, color=ACCENT)
    return paragraph


def _page_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def _page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    _font(run, size=8, color=TEXT_GRAY)


def _page_furniture(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(header.add_run("歐立威科技｜資料庫健檢報告"), size=8, color=TEXT_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(
        footer.add_run("(02) 2558-2656　｜　https://www.omniwaresoft.com.tw/　　"),
        size=8,
        color=TEXT_GRAY,
    )
    _page_field(footer)


def _configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    _page_furniture(section)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    for level, size in ((1, 16), (2, 12), (3, 10.5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(ACCENT)


def _cover(doc, model):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(model.product), size=27, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(model.customer), size=18, color=ACCENT)
    if model.system_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(f"系統／主機：{model.system_name}"), size=10.5, color=TEXT_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(f"{model.period} 資料庫健檢報告"), size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(f"Database Product：{model.product}"), size=10, color=TEXT_GRAY)
    bar = doc.add_table(rows=1, cols=1)
    _remove_borders(bar)
    _cell_text(bar.cell(0, 0), "歐立威科技", bold=True, color="FFFFFF")
    bar.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade(bar.cell(0, 0), TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(f"資料庫工程師 {model.engineer}"), size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("(02) 2558-2656"), size=9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("https://www.omniwaresoft.com.tw/"), size=9)


def _toc(doc):
    _heading(doc, "目錄")
    items = [
        ("1. 專案說明", "3"),
        ("　1.1 專案背景與目的", "3"),
        ("　1.2 維護週期與健檢頻率", "3"),
        ("2. 系統架構與環境", "4"),
        ("　2.1 架構總覽", "4"),
        ("3. 作業系統健檢", "5"),
        ("4. PostgreSQL 資料庫健檢", "依內容"),
        ("5. PEM 與 EFM", "依內容"),
        ("6. 更新與建議", "末章"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(title), size=10.5, bold=not title.startswith("　"))
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(16), alignment=2, leader=1)
        _font(p.add_run(f"\t{page}"), size=10)


def render_docx(model: ReportModel, output_path: Path) -> None:
    doc = Document()
    _configure(doc)
    _cover(doc, model)

    _page_break(doc)
    _toc(doc)

    _page_break(doc)
    _heading(doc, "1. 專案說明")
    _heading(doc, "1.1 專案背景與目的", 2)
    for text in (
        "1. 確認各節點、作業系統與資料庫運行狀態",
        "2. 檢查交易、權限、設定與容量風險",
        "3. 提出精簡且可執行的改善建議",
    ):
        doc.add_paragraph(text)
    _heading(doc, "1.2 維護週期與健檢頻率", 2)
    doc.add_paragraph(f"本次健檢期間：{model.period}；證據覆蓋率：{model.coverage['coverage_percent']}%。")

    _page_break(doc)
    _heading(doc, "2. 系統架構與環境")
    _heading(doc, "2.1 架構總覽", 2)
    _output_block(
        doc,
        ["節點", "角色", "服務"],
        [[node["hostname"], node["role"], "、".join(node["services"]) or "-"] for node in model.nodes],
    )
    _heading(doc, "2.2 健檢結果摘要", 2)
    _output_block(
        doc,
        ["正常", "注意", "嚴重", "待確認"],
        [[model.summary.get(key, 0) for key in ("normal", "attention", "critical", "pending")]],
    )

    for section in model.sections:
        _page_break(doc)
        _heading(doc, f"{section['section_id']}. {section['title']}")
        for group in section["groups"]:
            _heading(doc, group["title"], 2)
            for check in group["units"]:
                _heading(doc, check["title"], 3)
                _output_block(doc, check["headers"], check["rows"])
                if check["omitted_rows"]:
                    p = doc.add_paragraph(
                        f"報告顯示前 {len(check['rows'])} 筆；其餘 {check['omitted_rows']} 筆保留於內部證據。"
                    )
                    for run in p.runs:
                        _font(run, size=8, color=TEXT_GRAY)
                if check["assessment"]:
                    _assessment_block(doc, check["assessment"])

    _page_break(doc)
    _heading(doc, "6. 更新與建議")
    _heading(doc, "6.1 版本更新資訊摘要", 2)
    doc.add_paragraph(model.cve["message"])
    _heading(doc, "可修正 CVE 清單", 3)
    _output_block(
        doc,
        ["CVE", "修正版本", "CVSS", "修正內容", "權威來源"],
        [["待確認", "-", "-", "尚未接入 CVE 資料來源", "-"]],
    )
    _heading(doc, "6.2 健檢結論與優化建議", 2)
    for finding in model.findings:
        _heading(doc, f"{finding['title']} ({finding['node']})", 3)
        _assessment_block(
            doc,
            {
                "status": finding["status"],
                "status_label": finding["status_label"],
                "observation": finding["observation"],
                "recommendation": finding["recommendation"],
            },
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice is required for PDF conversion")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix=".lo-profile-", dir=pdf_path.parent) as profile:
        environment = os.environ.copy()
        environment["HOME"] = profile
        environment["XDG_CONFIG_HOME"] = str(Path(profile) / "xdg_config")
        environment["XDG_CACHE_HOME"] = str(Path(profile) / "xdg_cache")
        Path(environment["XDG_CONFIG_HOME"]).mkdir()
        Path(environment["XDG_CACHE_HOME"]).mkdir()
        environment["TMPDIR"] = "/private/tmp" if Path("/private/tmp").is_dir() else profile
        font_config = Path(__file__).parents[2] / "config" / "fonts.macos.conf"
        if font_config.is_file():
            environment["FONTCONFIG_FILE"] = str(font_config)
        result = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation={Path(profile).resolve().as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            check=False,
            env=environment,
        )
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if result.returncode or not generated.exists():
        raise RuntimeError(f"PDF conversion failed: {result.stderr.strip()}")
    if generated != pdf_path:
        generated.replace(pdf_path)
