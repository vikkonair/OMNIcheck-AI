#!/usr/bin/env python3
"""Build the living OMNIcheck AI operations guide DOCX from Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "OMNICHECK_AI_BUILD_AND_OPERATIONS_GUIDE.md"
OUTPUT = ROOT / "docs" / "OMNICHECK_AI_BUILD_AND_OPERATIONS_GUIDE.docx"
BLUE = "2E74B5"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F2F6FA"
GRAY = "666666"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=45, start=120, bottom=45, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr, fld_char2))


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "在 Word 中按 Ctrl+A、F9 更新目錄"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr, fld_char2, placeholder, fld_char3))


def set_fonts(run, latin="Calibri", east_asia="Microsoft JhengHei") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_inline(paragraph, text: str, *, code=False, bold=False, italic=False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        set_fonts(run, "Consolas", "Microsoft JhengHei")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    else:
        set_fonts(run)


INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|<https?://[^>]+>)")


def add_rich_text(paragraph, text: str) -> None:
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            add_inline(paragraph, text[position : match.start()])
        token = match.group(0)
        if token.startswith("`"):
            add_inline(paragraph, token[1:-1], code=True)
        elif token.startswith("**"):
            add_inline(paragraph, token[2:-2], bold=True)
        else:
            url = token[1:-1]
            add_inline(paragraph, url)
        position = match.end()
    if position < len(text):
        add_inline(paragraph, text[position:])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18

    heading_specs = {
        "Title": (27, NAVY, 0, 10),
        "Subtitle": (12, GRAY, 0, 8),
        "Heading 1": (16, BLUE, 18, 9),
        "Heading 2": (13, BLUE, 13, 6),
        "Heading 3": (11.5, NAVY, 9, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    code_style = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    code_style.font.size = Pt(8)
    code_style.font.color.rgb = RGBColor.from_string(NAVY)
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing = 1.02

    quote_style = document.styles.add_style("Guide Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = "Calibri"
    quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    quote_style.font.size = Pt(9.5)
    quote_style.font.color.rgb = RGBColor.from_string(NAVY)
    quote_style.paragraph_format.left_indent = Inches(0.28)
    quote_style.paragraph_format.right_indent = Inches(0.16)
    quote_style.paragraph_format.space_before = Pt(5)
    quote_style.paragraph_format.space_after = Pt(7)

    document.settings.element.append(OxmlElement("w:updateFields"))
    document.settings.element[-1].set(qn("w:val"), "true")


def add_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.text = "OMNIcheck AI  |  建置、部署與維運主手冊"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_fonts(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(footer, "Omniwaresoft Tech  ·  內部使用  ·  ")
    add_page_field(footer)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)


def add_cover(document: Document, lines: list[str]) -> int:
    title = lines[0].removeprefix("# ")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(74)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("OMNIcheck AI")
    set_fonts(run)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title_p = document.add_paragraph(style="Title")
    add_rich_text(title_p, title)
    subtitle = document.add_paragraph(style="Subtitle")
    add_rich_text(subtitle, "從空白 VM 到可驗證、可維運、可回復的完整系統")

    bar = document.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    bar.autofit = False
    bar.columns[0].width = Inches(6.75)
    cell = bar.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell, 80, 180, 80, 180)
    p = cell.paragraphs[0]
    r = p.add_run("BUILD  ·  DEPLOY  ·  OPERATE  ·  RECOVER")
    set_fonts(r)
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.bold = True
    r.font.size = Pt(10)

    index = 1
    while index < len(lines) and not lines[index].strip():
        index += 1
    meta = []
    while index < len(lines) and lines[index].strip():
        meta.append(lines[index].rstrip().rstrip("  "))
        index += 1
    document.add_paragraph()
    for item in meta:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_rich_text(p, item)
    document.add_paragraph()
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("Omniwaresoft Tech")
    set_fonts(r)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    document.add_page_break()
    toc_title = document.add_paragraph("目錄", style="Heading 1")
    set_keep_with_next(toc_title)
    toc_entries = [
        line.removeprefix("## ")
        for line in lines
        if line.startswith("## ") and not line.startswith("### ")
    ]
    midpoint = (len(toc_entries) + 1) // 2
    toc_table = document.add_table(rows=midpoint, cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    toc_table.autofit = False
    for position, entry in enumerate(toc_entries):
        column = 0 if position < midpoint else 1
        row = position if column == 0 else position - midpoint
        cell = toc_table.cell(row, column)
        set_cell_margins(cell, 45, 80, 45, 80)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(entry)
        set_fonts(run)
        run.font.size = Pt(9.2)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    document.add_page_break()
    return index + 1


def add_table(document: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    available = 6.86
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for col_index in range(width):
            cell = row.cells[col_index]
            cell.width = Inches(available / width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            text = values[col_index] if col_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_rich_text(paragraph, text)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
        if row_index == 0:
            set_repeat_table_header(row)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(document: Document, code_lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, 90, 130, 90, 130)
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["Code Block"]
    paragraph.paragraph_format.keep_together = False
    run = paragraph.add_run("\n".join(code_lines))
    set_fonts(run, "Consolas", "Microsoft JhengHei")
    run.font.size = Pt(8)


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    parts = split_table_row(line)
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", part) for part in parts)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    add_header_footer(document.sections[0])
    index = add_cover(document, lines)
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1]):
            table_rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, table_rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_rich_text(paragraph, text)
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Guide Quote")
            add_rich_text(paragraph, line[2:])
            index += 1
            continue
        numbered = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        bullet = re.match(r"^(\s*)-\s+(.+)$", line)
        if numbered:
            depth = 2 if len(numbered.group(1)) >= 2 else 1
            style = "List Number 2" if depth == 2 else "List Number"
            paragraph = document.add_paragraph(style=style)
            add_rich_text(paragraph, numbered.group(3))
            index += 1
            continue
        if bullet:
            depth = 2 if len(bullet.group(1)) >= 2 else 1
            style = "List Bullet 2" if depth == 2 else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            add_rich_text(paragraph, bullet.group(2))
            index += 1
            continue

        paragraph_parts = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].rstrip()
            if not candidate.strip():
                index += 1
                break
            if (
                candidate.startswith(("#", ">", "```", "|", "- "))
                or re.match(r"^\s*\d+\.\s+", candidate)
            ):
                break
            paragraph_parts.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_rich_text(paragraph, " ".join(paragraph_parts))

    core = document.core_properties
    core.title = "OMNIcheck AI 建置、部署與維運主手冊"
    core.subject = "可重建、部署、驗證、維運與回復的系統手冊"
    core.author = "Omniwaresoft Tech"
    core.keywords = "OMNIcheck AI, EDB, EPAS, deployment, operations, runbook"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
